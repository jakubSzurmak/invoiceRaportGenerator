from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


def load_data(paid, issued, cash):

    document = Document()
    title = document.add_paragraph()
    styled_title = title.add_run('Zestawienie faktur za X XXXX r.')
    styled_title.bold = True
    styled_title.font.size = Pt(14)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph('')
    document.add_paragraph('')

    if len(paid) <= 0 or not paid:
        pass
    else:
        paid_title = document.add_paragraph()
        paid_title_styled = paid_title.add_run('Faktury kosztowe:')
        paid_title_styled.bold = True
        paid_title_styled.underline = True
        paid_title_styled.font.size = Pt(12)
        for i in range(len(paid)):
            base = document.add_paragraph()
            prepare_lines(paid, i, base)

    if len(issued) <= 0 or not issued:
        pass
    else:
        issued_title = document.add_paragraph()
        issued_title_styled = issued_title.add_run('Faktury wystawione:')
        issued_title_styled.bold = True
        issued_title_styled.underline = True
        issued_title_styled.font.size = Pt(12)
        for j in range(len(issued)):
            base = document.add_paragraph()
            prepare_lines(issued, j, base)

    if len(cash) <= 0 or not cash:
        pass
    else:
        cash_title = document.add_paragraph()
        cash_title_styled = cash_title.add_run('Wynagrodzenia:')
        cash_title_styled.bold = True
        cash_title_styled.underline = True
        cash_title_styled.font.size = Pt(12)
        for k in range(len(issued)):
            base = document.add_paragraph()
            prepare_lines(cash, k, base)

    document.save('Zestawienie faktur za X')


def prepare_lines(x_lines, ind, base_p):
    done_line = x_lines[ind]
    styled_base = base_p.add_run(done_line)
    styled_base.font.size = Pt(11)


